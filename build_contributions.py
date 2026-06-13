# -*- coding: utf-8 -*-
"""Populate 'COS40007 - Who did What.docx' with each team member's AI contribution."""
import os
import docx
from docx.shared import Pt

ROOT = os.path.dirname(os.path.abspath(__file__))
DOC  = os.path.join(ROOT, "COS40007 - Who did What.docx")

CONTRIBUTIONS = {
    "Kenneth Hui Hong CHUA": [
        "AI Role: Time-Series Modelling Lead (ARIMA).",
        "• Sourced and preprocessed the national CPI inflation dataset from OpenDOSM "
        "(cpi_2d_inflation, 7,770 monthly records, 14 divisions, 1980-2026).",
        "• Implemented the ARIMA module (src/time_series.py): BIC grid search over "
        "p in [0,4] and q in [0,4] for order selection.",
        "• Performed ADF stationarity test on the training series to determine the "
        "differencing order d.",
        "• Trained the final ARIMA model on the complete series and computed AIC and BIC "
        "for model justification.",
        "• Computed time-series evaluation metrics on the hold-out test set (MAE, RMSE, MAPE) "
        "and analysed forecast accuracy.",
    ],
    "Kelvin Wen Kiong FONG": [
        "AI Role: State Clustering Lead.",
        "• Sourced and preprocessed the state CPI and household income datasets "
        "(cpi_2d_state, hh_income_state) covering 16 Malaysian states.",
        "• Engineered the 5 state-level features (mean CPI index, CPI growth rate, "
        "CPI volatility, latest mean income, latest median income) and applied z-score "
        "standardisation.",
        "• Implemented the clustering module (src/clustering.py): K-Means with elbow + "
        "silhouette + Davies-Bouldin selection, and Agglomerative Hierarchical with Ward linkage.",
        "• Performed PCA dimensionality reduction (2 components) for cluster visualisation and "
        "assigned interpretable semantic labels to each cluster.",
        "• Computed the three internal validity indices (silhouette, Davies-Bouldin, "
        "Calinski-Harabasz) to validate cluster structure across both algorithms.",
        "• Produced cluster_pca_kmeans.png, cluster_pca_hierarchical.png, cluster_dendrogram.png, "
        "cluster_elbow_silhouette.png, and cluster_profiles.png.",
    ],
    "Nigel Zi Jun LING": [
        "AI Role: Cross-Model Evaluation & Analytics.",
        "• Sourced and preprocessed the retail fuel price dataset (fuelprice) and resampled "
        "weekly data to monthly frequency to align with the CPI series.",
        "• Implemented fuel-price-to-CPI pass-through analysis in src/evaluation.py: "
        "contemporaneous and lag-1 Pearson correlations for RON95, RON97, and diesel.",
        "• Built the cross-state CPI comparison and the 13-division x year inflation heatmap "
        "to surface category-level inflation hotspots.",
        "• Generated the income-vs-CPI scatter coloured by cluster to validate the link "
        "between the time-series and clustering results.",
        "• Cross-validated every headline statistic against official DOSM 2022 published "
        "figures (Sabah mean income, Selangor mean income, national inflation trend).",
    ],
    "Neng Yi CHIENG": [
        "AI Role: Time-Series Forecasting Co-Lead + Pipeline / Demonstrator Architect.",
        "• Co-developed the ARIMA forecasting workflow: train/test split logic, refitting on the "
        "full series for the out-of-sample forecast, and 95% confidence-interval extraction.",
        "• Built the time-series visualisation suite (src/time_series.py plotting functions): "
        "ts_forecast.png (train + test + forecast with CI band) and ts_test_evaluation.png "
        "(actual vs predicted on the hold-out).",
        "• Built the 24-month out-of-sample inflation forecast generator and the JSON serialisation "
        "of all ARIMA artefacts (arima_results.json, ts_metrics_summary.csv).",
        "• Implemented division-level inflation trend analysis (ts_division_trends.png) to support "
        "the forecast interpretation.",
        "• Designed and implemented the end-to-end reproducible AI pipeline "
        "(run_pipeline.py, run_analysis.py) that orchestrates preprocessing -> ARIMA -> "
        "clustering -> evaluation deterministically with random_state = 42.",
        "• Built the Flask web dashboard (app/dashboard.py) at http://localhost:5050 with 4 tabs "
        "(Overview, Time-Series Forecast, Clustering, Fuel Analysis).",
        "• Implemented the JSON REST API layer (/api/overview, /api/ts/forecast, "
        "/api/cluster/states, /api/fuel/correlation) that serves model artefacts to the UI.",
        "• Integrated Git LFS for the raw datasets and managed the project repository structure.",
    ],
}


def main():
    doc = docx.Document(DOC)
    # The second table (index 1) has rows: header + 4 member rows with empty second column
    target = doc.tables[1]
    name_to_row = {}
    for row in target.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2 and cells[0] in CONTRIBUTIONS:
            name_to_row[cells[0]] = row

    if len(name_to_row) != 4:
        raise RuntimeError(f"Expected 4 member rows, found {len(name_to_row)}: {list(name_to_row)}")

    for name, row in name_to_row.items():
        cell = row.cells[1]
        # Clear and write fresh content
        cell.text = ""
        first = True
        for line in CONTRIBUTIONS[name]:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)
            if line.startswith("AI Role:"):
                run.bold = True
            first = False

    doc.save(DOC)
    print("Contributions written to:", DOC)
    for name in CONTRIBUTIONS:
        print(f"  - {name}: {len(CONTRIBUTIONS[name])} lines")


if __name__ == "__main__":
    main()
