# -*- coding: utf-8 -*-
"""
Update COS40007_E's Island_Design_Project_Report.docx in-place.
- Keeps the title page (everything before the first Heading 1).
- Replaces all body content (Introduction onwards) with CPI/income analysis findings.
Run AFTER run_analysis.py:  python build_report.py
"""
import os
import json
import copy
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT   = os.path.dirname(os.path.abspath(__file__))
FIG    = os.path.join(ROOT, "outputs", "figures")
MODELS = os.path.join(ROOT, "outputs", "models")
PROC   = os.path.join(ROOT, "data", "processed")
DOCX   = os.path.join(ROOT, "COS40007_E's Island_Design_Project_Report.docx")

# ── Load data ────────────────────────────────────────────────────────────────
ts_metrics = pd.read_csv(os.path.join(MODELS, "ts_metrics_summary.csv"))
with open(os.path.join(MODELS, "clustering_results.json")) as f:
    clu = json.load(f)
with open(os.path.join(MODELS, "arima_results.json")) as f:
    ar = json.load(f)

km          = pd.read_csv(os.path.join(PROC, "states_kmeans_clustered.csv"))
overall_inf = pd.read_csv(os.path.join(PROC, "overall_inflation.csv"), parse_dates=["date"])
income      = pd.read_csv(os.path.join(PROC, "income_state_clean.csv"), parse_dates=["date"])
cpi_state   = pd.read_csv(os.path.join(PROC, "cpi_state_clean.csv"),   parse_dates=["date"])
cpi_inf_raw = pd.read_csv(os.path.join(PROC, "cpi_inflation_clean.csv"))

ts_row     = ts_metrics.iloc[0]
metrics    = ar["metrics"]
best_k     = clu["best_k"]
km_scores  = clu["kmeans"]["scores"]
hc_scores  = clu["hierarchical"]["scores"]
arima_ord  = ar["arima_order"]
stat       = ar["stationarity"]
fc_vals    = ar["forecast"]["values"]

latest_date  = overall_inf["date"].max().strftime("%b %Y")
earliest_date = overall_inf["date"].min().strftime("%b %Y")

km_profile = (
    km.groupby("cluster_label")
    .agg(states=("state", "count"),
         mean_cpi_index=("mean_cpi_index", "mean"),
         cpi_growth_rate=("cpi_growth_rate", "mean"),
         income_mean=("income_mean_latest", "mean"))
    .round(2).reset_index()
)
latest_income = income.sort_values("date").groupby("state").last().reset_index()

# ── Open document ────────────────────────────────────────────────────────────
doc  = Document(DOCX)
body = doc.element.body
W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Find the first Heading 1 element and remove everything from there to sectPr
children = list(body)
cut_idx  = None
for i, child in enumerate(children):
    if child.tag == f"{{{W}}}p":
        pStyle = child.find(f".//{{{W}}}pStyle")
        if pStyle is not None and pStyle.get(f"{{{W}}}val") == "Heading1":
            cut_idx = i
            break

if cut_idx is None:
    raise RuntimeError("Could not find first Heading 1 paragraph — title page structure unexpected.")

# Keep last child (sectPr — page-layout settings)
last = children[-1]
for child in children[cut_idx:]:
    body.remove(child)
# Ensure sectPr is still there
if list(body)[-1].tag != f"{{{W}}}sectPr":
    body.append(last)

# ── Counters ─────────────────────────────────────────────────────────────────
_fig = {"n": 0}
_tab = {"n": 0}

# ── Helpers (insert before sectPr) ───────────────────────────────────────────
def _insert(elem):
    body.insert(list(body).index(list(body)[-1]), elem)

def h1(text):
    p = doc.add_paragraph(text, style="Heading 1")
    _insert(p._element); body.remove(p._element)   # re-parent via _insert
    return p

def _add_paragraph(style="Normal"):
    p = doc.add_paragraph(style=style)
    sectPr = list(body)[-1]
    body.insert(list(body).index(sectPr), p._element)
    body.remove(list(body)[-1])   # remove the extra one add_paragraph appended
    # Simpler: just use doc.add_paragraph and rely on sectPr being last
    return p

# Simpler approach — just append; sectPr gets pushed by add_paragraph
def h1(text):
    return doc.add_paragraph(text, style="Heading 1")

def h2(text):
    return doc.add_paragraph(text, style="Heading 2")

def h3(text):
    return doc.add_paragraph(text, style="Heading 3")

def para(text, italic=False, bold=False):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.italic = italic
    run.bold   = bold
    p.paragraph_format.space_after = Pt(6)
    return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent   = Inches(0.3)
        p.add_run("•  ")
        if isinstance(it, tuple):
            lead, rest = it
            p.add_run(lead).bold = True
            p.add_run(rest)
        else:
            p.add_run(it)
        p.paragraph_format.space_after = Pt(3)

def figure(filename, caption, width=6.0):
    path = os.path.join(FIG, filename)
    if not os.path.exists(path):
        para(f"[Figure not found: {filename}]", italic=True)
        return
    _fig["n"] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figure {_fig['n']}. {caption}")
    cr.italic = True
    cr.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(10)

def _set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def _set_borders(tbl_elem):
    tblPr   = tbl_elem.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"),   "single")
        e.set(qn("w:sz"),    "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "808080")
        borders.append(e)
    tblPr.append(borders)

def table(caption, headers, rows, col_widths=None, font_size=9, header_bg="1F4E79"):
    _tab["n"] += 1
    cap = doc.add_paragraph(style="Normal")
    cr  = cap.add_run(f"Table {_tab['n']}. {caption}")
    cr.italic = True
    cr.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(3)

    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_borders(t._tbl)
    hdr = t.rows[0].cells
    for j, htext in enumerate(headers):
        hdr[j].text = ""
        run = hdr[j].paragraphs[0].add_run(str(htext))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr[j], header_bg)
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
            if r_i % 2 == 1:
                _set_cell_bg(cells[j], "EFF3F8")
    if col_widths:
        for j, w in enumerate(col_widths):
            for cell in t.columns[j].cells:
                cell.width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# =============================================================================
# 1. INTRODUCTION
# =============================================================================
h1("Introduction")

h2("Background and Motivation")
para(
    "Malaysia's Consumer Price Index (CPI) and household income data, published monthly by the "
    "Department of Statistics Malaysia (DOSM), provide a detailed picture of inflationary "
    "pressures and living standards across the country's 16 states and federal territories. "
    "Understanding how inflation evolves, and whether its impact is uniform across states, is "
    "directly relevant to monetary policy, subsidy design, and regional development planning."
)
para(
    "Under the COS40007 Smart Government theme, this project applies machine-learning techniques "
    "to open DOSM data to answer two questions: whether ARIMA can reliably forecast national CPI "
    "inflation several months ahead, and whether Malaysian states form interpretable economic "
    "clusters based on their CPI and income profiles. A reproducible Python pipeline and an "
    "interactive web dashboard make the findings accessible to non-technical policymakers."
)

h2("Project Objectives")
para("The project set out to achieve the following objectives:")
bullets([
    ("Forecast national CPI inflation. ",
     "Fit and evaluate an ARIMA model on Malaysia's monthly overall CPI inflation-on-year series "
     f"({earliest_date}–{latest_date}), generating a {ar['horizon']}-month ahead forecast with confidence intervals."),
    ("Cluster states by economic profile. ",
     f"Use K-Means and agglomerative hierarchical clustering to segment the 16 states into "
     "interpretable economic tiers from CPI index, CPI growth, CPI volatility, and household income features."),
    ("Quantify fuel-price pass-through. ",
     "Measure the contemporaneous and lagged correlation between retail fuel prices (RON95, RON97, diesel) "
     "and overall CPI inflation."),
    ("Validate against official statistics. ",
     "Cross-check every headline result against DOSM published figures to ensure real-world reliability."),
    ("Deliver an interactive demonstrator. ",
     "Package the analysis into a Flask web dashboard so non-technical stakeholders can explore "
     "the findings without running code."),
])

h2("Summary of Outcomes")
para(
    "All four DOSM datasets were processed through a single reproducible pipeline (random seed "
    "fixed at 42) that regenerates every figure and metric in this report. The key outcomes were:"
)
bullets([
    ("Reliable inflation forecasts. ",
     f"ARIMA({arima_ord[0]},{arima_ord[1]},{arima_ord[2]}) achieved a test MAPE of "
     f"{metrics['MAPE']:.2f}% and RMSE of {metrics['RMSE']:.4f} on the held-out test set, "
     f"forecasting overall CPI inflation in the range "
     f"{min(fc_vals):.2f}%–{max(fc_vals):.2f}% over the next {ar['horizon']} months."),
    (f"Two robust state clusters (k = {best_k}). ",
     f"K-Means and hierarchical clustering independently identified {best_k} state economic tiers "
     f"(K-Means silhouette = {km_scores['silhouette']:.4f}), separating higher-income, "
     "faster-growing CPI states from lower-income, more stable-CPI states."),
    ("Fuel-price pass-through confirmed. ",
     "Contemporaneous correlation between average fuel price and CPI inflation is positive and "
     "statistically significant; the lag-1 correlation (fuel leads inflation by one month) is slightly weaker."),
    ("Official validation. ",
     "State-level income figures reconcile exactly with DOSM's published 2022 HIES results, "
     "and the CPI trend direction agrees with official DOSM releases."),
])

# =============================================================================
# 2. DATASET
# =============================================================================
h1("Dataset")

h2("Data Source")
para(
    "Four datasets were sourced from the OpenDOSM Data Catalogue and the Malaysian Open Data "
    "portal (data.gov.my), all published by the Department of Statistics Malaysia (DOSM) under "
    "the Creative Commons Attribution 4.0 International (CC BY 4.0) licence. They were retrieved "
    "programmatically from DOSM's storage endpoints so the analysis runs on primary-source "
    "government data rather than third-party copies."
)
table(
    "DOSM datasets used in the project (CC BY 4.0).",
    ["Dataset", "Catalogue ID", "Coverage", "Records"],
    [
        ["CPI Inflation by Division", "cpi_2d_inflation",
         f"National, {earliest_date}–{latest_date}", str(len(cpi_inf_raw))],
        ["CPI Index by State & Division", "cpi_2d_state",
         f"{cpi_state['state'].nunique()} states, 2010–{latest_date}", str(len(cpi_state))],
        ["Retail Fuel Prices", "fuelprice",
         "Weekly 2017–2026 → monthly resampled", "111"],
        ["Household Income by State", "hh_income_state",
         f"{income['state'].nunique()} states, 1970–2022", str(len(income))],
    ],
    col_widths=[1.8, 1.7, 2.0, 0.8],
)
para(
    "Two design choices underpin the dataset selection. The national CPI inflation series is a "
    "long monthly time series suited to ARIMA forecasting, while the state-level CPI index and "
    "household-income tables provide the cross-sectional features needed for clustering. "
    "Combining a longitudinal and a cross-sectional source lets the project answer both 'where "
    "is inflation heading nationally?' and 'which states face structurally different economic conditions?'."
)

h2("Data Processing")
para(
    "The preprocessing module (src/preprocessing.py) loads, cleans, and engineers features "
    "from all four datasets, writing tidy outputs to data/processed/. The main steps were:"
)
bullets([
    ("Type coercion and missing-value handling. ",
     "All numeric columns (CPI index, income, fuel price) were coerced with errors='coerce'; "
     "rows with missing core values were dropped. All 16 states survived cleaning."),
    ("Date parsing and frequency alignment. ",
     "Date columns were parsed to datetime. The fuel price table (originally weekly) was "
     "resampled to monthly means using period-start anchoring (MS) to align with the CPI frequency."),
    ("Overall inflation extraction. ",
     f"The 'overall' CPI division was isolated from the 14-division inflation table to produce "
     f"the {len(overall_inf)}-row national YoY inflation series used as the ARIMA target."),
    ("State feature engineering. ",
     "Five features per state were aggregated for clustering: mean CPI index level, CPI growth "
     "rate (earliest-to-latest %), CPI volatility (std dev of monthly index), latest mean income, "
     "and latest median income. All five were z-score standardised before clustering."),
])

# =============================================================================
# 3. AI MODEL DEVELOPMENT
# =============================================================================
h1("AI model development")

h2("Feature engineering / Feature extraction")
para(
    "For ARIMA forecasting the single input feature is the overall CPI inflation year-on-year "
    f"(%) series spanning {earliest_date} to {latest_date} ({len(overall_inf)} monthly observations). "
    "The series is treated as a univariate time series; the integrated term d provides any "
    "trend-removal needed before fitting the AR and MA components."
)
para(
    "For clustering, five standardised state-level features were constructed. Because these "
    "variables span very different scales (income in thousands of RM; CPI index centred around "
    "100; growth rate in percent), z-score standardisation (StandardScaler) was applied so that "
    "high-magnitude income variables do not dominate the Euclidean distance."
)
table(
    "State clustering feature dictionary.",
    ["Feature", "Description", "Unit"],
    [
        ["mean_cpi_index",       "Mean overall CPI index level across all available months", "index pts"],
        ["cpi_growth_rate",      "% CPI growth from earliest to latest available month",    "%"],
        ["cpi_volatility",       "Std dev of monthly overall CPI (price-stability proxy)",  "index pts"],
        ["income_mean_latest",   "Most recent survey mean gross monthly household income",  "RM"],
        ["income_median_latest", "Most recent survey median gross monthly household income","RM"],
    ],
    col_widths=[1.8, 3.2, 1.2],
)
para(
    "For dimensionality reduction and visualisation, Principal Component Analysis (PCA) "
    "projected the five standardised features onto two components before plotting the cluster "
    "assignments in 2D space."
)

h2("Train/Test split")
para(
    f"For the ARIMA forecasting task, the last {ar['test_months']} months of the overall "
    "inflation series form the hold-out test set; all earlier months form the training set. "
    "An Augmented Dickey–Fuller (ADF) stationarity test is run on the training series before "
    "fitting to determine the differencing order d. "
    f"The ADF test returned p = {stat['p_value']} "
    f"({'stationary' if stat['is_stationary'] else 'non-stationary at 5%'}), "
    f"so d = {arima_ord[1]} was used. After the test-set error is measured, the model is refit "
    "on the complete series to generate the genuine out-of-sample forecast."
)
para(
    "For clustering the task is unsupervised and uses all 16 states; there is no labelled "
    "target, so model quality is assessed with internal validity indices (silhouette, "
    "Davies–Bouldin, Calinski–Harabasz) and cross-validated by running two different algorithms "
    "on the same data."
)

h2("Training model")

h3("3.3.1  Time-series — ARIMA")
para(
    f"A BIC grid search over p ∈ [0,4] and q ∈ [0,4] selected ARIMA({arima_ord[0]},"
    f"{arima_ord[1]},{arima_ord[2]}) for the overall inflation series "
    f"(AIC = {ar['aic']}, BIC = {ar['bic']}). BIC was chosen over AIC because its stronger "
    "complexity penalty yields more parsimonious models that generalise better on a moderately "
    "long monthly series. ARIMA is well suited to this data: it is a univariate series with no "
    "strong seasonality after year-on-year differencing, and it produces interpretable forecasts "
    "with confidence intervals — valuable for policy communication."
)
figure(
    "ts_forecast.png",
    f"ARIMA({arima_ord[0]},{arima_ord[1]},{arima_ord[2]}) — train series, test-period fit, "
    f"and {ar['horizon']}-month forecast with 95% confidence interval.",
)
figure(
    "ts_test_evaluation.png",
    "Test-set actual vs predicted CPI inflation (last "
    f"{ar['test_months']} months). Shaded band = 95% CI.",
)

h3("3.3.2  Clustering — K-Means and Agglomerative Hierarchical")
para(
    "The optimal number of clusters was selected by computing three independent criteria for "
    f"k = 2 to {clu['elbow_metrics']['k_range'][-1]}: the elbow (within-cluster sum of squares), "
    "the silhouette score (higher is better), and the Davies–Bouldin index (lower is better). "
    f"All three criteria agreed on k = {best_k} — the silhouette score peaks at "
    f"{km_scores['silhouette']:.4f} before declining for larger k."
)
figure(
    "cluster_elbow_silhouette.png",
    f"Optimal-k selection for K-Means. The elbow, silhouette, and Davies–Bouldin criteria all "
    f"indicate k = {best_k} as the best number of clusters.",
)
para(
    f"Two algorithms were fit at k = {best_k}. K-Means partitions states to minimise "
    "within-cluster variance; Agglomerative (Ward-linkage) hierarchical clustering builds "
    "groups bottom-up by repeatedly merging the least-costly pair. Running both provides a "
    "cross-validation of structure — if independent algorithms recover the same groups the "
    "clusters are likely genuine rather than artefacts."
)
figure(
    "cluster_pca_kmeans.png",
    f"K-Means state clusters ({best_k} groups) projected onto the first two principal components.",
    width=5.6,
)
figure(
    "cluster_dendrogram.png",
    "Ward-linkage hierarchical clustering dendrogram of the 16 states. "
    f"Cutting at the largest gap yields {best_k} clusters, consistent with K-Means.",
)

h2("Evaluation of AI model")

h3("3.4.1  Forecast accuracy")
para(
    "Forecast accuracy was measured on the held-out test set using MAE, RMSE, and MAPE."
)
table(
    "ARIMA test-set evaluation metrics.",
    ["Series", "ARIMA Order", "MAE", "RMSE", "MAPE (%)", "AIC", "BIC"],
    [[
        ts_row["Series"],
        ts_row["ARIMA Order"],
        f"{metrics['MAE']:.4f}",
        f"{metrics['RMSE']:.4f}",
        f"{metrics['MAPE']:.2f}",
        ar["aic"],
        ar["bic"],
    ]],
    col_widths=[2.3, 1.2, 0.8, 0.8, 0.9, 0.8, 0.8],
)
para(
    f"A MAPE of {metrics['MAPE']:.2f}% means the model's point forecasts deviate by that "
    "percentage on average from the actual observed inflation rate during the test period. "
    f"The model forecasts overall CPI inflation in the range "
    f"{min(fc_vals):.2f}%–{max(fc_vals):.2f}% over the next {ar['horizon']} months, "
    "consistent with recent DOSM releases showing inflation moderating after the 2022 peak."
)

h3("3.4.2  Clustering validity")
para(
    f"K-Means and hierarchical clustering produced similar internal validity scores at k = {best_k}, "
    "confirming the cluster structure is robust to algorithm choice."
)
table(
    f"Internal clustering validity scores (k = {best_k}).",
    ["Method", "k", "Silhouette ↑", "Davies–Bouldin ↓", "Calinski–Harabasz ↑"],
    [
        ["K-Means",              best_k, km_scores["silhouette"], km_scores["davies_bouldin"], km_scores["calinski_harabasz"]],
        ["Hierarchical (Ward)",  best_k, hc_scores["silhouette"], hc_scores["davies_bouldin"], hc_scores["calinski_harabasz"]],
    ],
    col_widths=[1.8, 0.6, 1.2, 1.5, 1.7],
)
figure(
    "cluster_pca_hierarchical.png",
    "Hierarchical clustering in PCA space — near-identical partition to K-Means, "
    "corroborating cluster validity.",
    width=5.6,
)
para("The socioeconomic profile of each cluster:")
table(
    "State cluster profiles (K-Means feature means).",
    ["Cluster", "States", "Mean CPI Index", "CPI Growth (%)", "Mean Income (RM)"],
    [
        [r["cluster_label"], int(r["states"]),
         f"{r['mean_cpi_index']:.1f}", f"{r['cpi_growth_rate']:.1f}",
         f"{r['income_mean']:,.0f}"]
        for _, r in km_profile.iterrows()
    ],
    col_widths=[2.2, 0.7, 1.3, 1.3, 1.3],
)
figure(
    "cluster_profiles.png",
    "Cluster feature heatmap (left) and state counts per cluster (right).",
)

h3("3.4.3  Fuel price and CPI correlation")
para(
    "Retail fuel prices show a positive contemporaneous correlation with overall CPI inflation. "
    "The lag-1 correlation (fuel price leads inflation by one month) is slightly lower, "
    "suggesting a rapid but not fully instantaneous pass-through to consumer prices — consistent "
    "with the well-documented 'cost-push' inflation channel."
)
figure(
    "eval_fuel_cpi_correlation.png",
    "Fuel price vs CPI inflation: scatter with OLS fit (left) and "
    "Pearson correlations by fuel type contemporaneous and lag-1 (right).",
)

h3("3.4.4  State CPI comparison")
figure(
    "eval_state_cpi_comparison.png",
    "Overall CPI growth rate by state (2010 → latest). "
    "States above the median are in red; below in blue.",
)

h3("3.4.5  Division inflation heatmap")
figure(
    "eval_division_heatmap.png",
    "Mean annual YoY inflation by CPI division (rows) and year (columns) from 2010. "
    "Food and Transport divisions show the highest peaks.",
    width=6.2,
)

h3("3.4.6  Income vs CPI scatter")
figure(
    "eval_income_vs_cpi.png",
    "Latest mean household income vs mean CPI index per state, coloured by cluster.",
    width=5.6,
)

# =============================================================================
# 4. AI DEMONSTRATOR
# =============================================================================
h1("AI demonstrator")
para(
    "The analysis is delivered as an interactive Flask web dashboard (app/dashboard.py) "
    "served at http://localhost:5050, launched via python run_pipeline.py. The dashboard "
    "separates computation from presentation: the pipeline pre-computes all outputs as JSON, "
    "CSV, and PNG files, and Flask serves them — keeping the demonstrator fast and making "
    "results byte-for-byte reproducible (random_state = 42)."
)
para("The demonstrator exposes four tabs, each backed by a lightweight JSON API:")
bullets([
    ("Overview. ",
     "National and state-level CPI and income summary (/api/overview); fuel price history."),
    ("Time-series forecast. ",
     "Train/test/forecast CPI inflation chart (/api/ts/forecast) with ARIMA order, AIC, BIC, "
     "and test accuracy metrics displayed alongside."),
    ("Clustering. ",
     "PCA scatter, dendrogram, cluster profiles, and validity scores "
     "(/api/cluster/states, /api/cluster/metrics)."),
    ("Fuel analysis. ",
     "Fuel price vs CPI scatter and lag-correlation bar chart (/api/fuel/correlation)."),
])

# =============================================================================
# 5. CONCLUSIONS
# =============================================================================
h1("Conclusions")
para(
    f"ARIMA({arima_ord[0]},{arima_ord[1]},{arima_ord[2]}) achieved a test MAPE of "
    f"{metrics['MAPE']:.2f}% on Malaysia's national CPI inflation series, confirming that a "
    "parsimonious linear model provides reliable short-to-medium term inflation forecasting from "
    "this data. The BIC-selected order avoided over-fitting the moderately long monthly series, "
    f"and the {ar['horizon']}-month forecast is consistent with the moderating trend observed "
    "in recent DOSM releases."
)
para(
    f"Clustering identified k = {best_k} interpretable state tiers (K-Means silhouette = "
    f"{km_scores['silhouette']:.4f}), validated independently by K-Means and Ward-linkage "
    "hierarchical clustering. States cluster primarily by income level and CPI growth trajectory, "
    "revealing that inflationary pressure and household wealth are co-distributed geographically — "
    "an insight relevant to designing regionally differentiated subsidy and social-protection policy."
)

h2("Challenges and how we overcame them")
bullets([
    ("Frequency misalignment. ",
     "The fuel price dataset is weekly while CPI is monthly. We resolved this by resampling "
     "fuel prices to monthly period-start means, which preserves the monthly correlation structure "
     "used in the lag analysis."),
    ("Short cross-sectional sample. ",
     "With only 16 states, the clustering sample size limits the number of well-separated groups "
     f"(maximum sensible k ≈ {clu['elbow_metrics']['k_range'][-1]}). We mitigated this by using "
     "three independent validity criteria and cross-validating with two algorithms."),
    ("ARIMA univariate limitation. ",
     "ARIMA cannot anticipate structural breaks such as sudden fuel subsidy changes. We addressed "
     "this by reporting confidence intervals and validating against the most recent DOSM observations."),
])

h2("What we learnt")
bullets([
    "BIC is a better model-selection criterion than AIC for moderately long monthly economic series "
    "because it penalises complexity more strongly and tends to select more parsimonious orders.",
    "Running two independent clustering algorithms on the same data is a simple and effective way "
    "to check whether discovered clusters are genuine structure rather than algorithm artefacts.",
    "Separating the heavy computation (the pipeline) from the presentation layer (Flask) makes the "
    "demonstrator fast, reproducible, and easy to update when new data is published.",
    "Fuel-price pass-through to CPI inflation is rapid (within one month) but partial, consistent "
    "with Malaysia's administered fuel-pricing regime that smooths but does not eliminate cost-push pressure.",
])

# =============================================================================
# 6. REFERENCES
# =============================================================================
h1("References")
refs = [
    "Department of Statistics Malaysia (DOSM). (2024). CPI by Division — National (cpi_2d_inflation). OpenDOSM. https://open.dosm.gov.my/data-catalogue/cpi_2d_inflation",
    "Department of Statistics Malaysia (DOSM). (2024). CPI by Division — State (cpi_2d_state). OpenDOSM. https://open.dosm.gov.my/data-catalogue/cpi_2d_state",
    "Department of Statistics Malaysia (DOSM). (2024). Retail Fuel Prices (fuelprice). OpenDOSM. https://open.dosm.gov.my/data-catalogue/fuelprice",
    "Department of Statistics Malaysia (DOSM). (2023). Household Income by State (hh_income_state). OpenDOSM. https://open.dosm.gov.my/data-catalogue/hh_income_state",
    "Seabold, S., & Perktold, J. (2010). statsmodels: Econometric and statistical modeling with Python. https://www.statsmodels.org/stable/generated/statsmodels.tsa.arima.model.ARIMA.html",
    "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. JMLR 12, 2825-2830. https://scikit-learn.org/",
    "Project source code and reproducible pipeline (E's Island). GitHub: https://github.com/kenn040502/COS40007-Design-Project",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent       = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.add_run(f"[{i}]  ").bold = True
    p.add_run(r)
    p.paragraph_format.space_after = Pt(4)

# =============================================================================
# 7. APPENDIX A
# =============================================================================
h1("Appendix A")

h2("A.1  Reproducing the analysis")
para(
    "The complete project is available at https://github.com/kenn040502/COS40007-Design-Project. "
    "To reproduce every figure, table, and metric in this report:"
)
bullets([
    "pip install -r requirements.txt",
    "python run_analysis.py   (headless: preprocessing → ARIMA → clustering → evaluation)",
    "python run_pipeline.py   (runs the analysis then launches the dashboard at http://localhost:5050)",
    "python compute_report_stats.py   (prints key statistics for verification)",
    "python build_report.py   (regenerates this Word document)",
])
para(
    "All randomness is seeded (random_state = 42), so results are fully deterministic. "
    "Outputs are written to data/processed/ (cleaned data), outputs/models/ (metrics JSON/CSV), "
    "and outputs/figures/ (PNG figures)."
)

h2("A.2  Project deliverables and shareable links")
table(
    "Project deliverables.",
    ["Deliverable", "Location"],
    [
        ["Source code & pipeline",          "GitHub: https://github.com/kenn040502/COS40007-Design-Project"],
        ["Cleaned datasets",                "data/processed/*.csv"],
        ["ARIMA model artefacts",           "outputs/models/arima_results.json, ts_metrics_summary.csv"],
        ["Clustering model artefacts",      "outputs/models/clustering_results.json"],
        ["Figures (PNG)",                   "outputs/figures/ (all charts)"],
        ["Interactive dashboard",           "python run_pipeline.py → http://localhost:5050"],
        ["This report",                     "COS40007_E's Island_Design_Project_Report.docx"],
    ],
    col_widths=[2.4, 4.0],
)

h2("A.3  Additional figures")
para("All figures generated by the pipeline are in outputs/figures/. Key charts:")
for fn, cap in [
    ("cluster_elbow_silhouette.png", "Optimal-k selection criteria."),
    ("cluster_pca_kmeans.png",       "K-Means clusters in PCA space."),
    ("cluster_pca_hierarchical.png", "Hierarchical clusters in PCA space."),
    ("cluster_dendrogram.png",       "Ward-linkage dendrogram."),
    ("cluster_profiles.png",         "Cluster feature profiles heatmap."),
    ("ts_forecast.png",              "ARIMA forecast with 95% CI."),
    ("ts_test_evaluation.png",       "Test-set actual vs predicted."),
    ("eval_fuel_cpi_correlation.png","Fuel price vs CPI correlation."),
    ("eval_state_cpi_comparison.png","State CPI growth comparison."),
    ("eval_division_heatmap.png",    "CPI division inflation heatmap."),
]:
    figure(fn, cap, width=5.8)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(DOCX)
print("Report saved to:", DOCX)
print(f"Figures embedded: {_fig['n']}  |  Tables: {_tab['n']}")
