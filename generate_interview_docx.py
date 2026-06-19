"""
Generate a detailed Word document for the lecturer interview preparation.
Split by team member with Q&A and code line explanations.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ─── Styles ──────────────────────────────────────────────────────────────────
styles = doc.styles

# Base font
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def add_title(text, size=22, color=(26, 60, 110)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p


def add_subtitle(text, size=14, color=(80, 80, 80)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p


def add_section_header(text, color=(26, 60, 110)):
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(*color)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1A3C6E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_subheader(text, color=(232, 82, 58), size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p


def add_role_box(role_text, files_text):
    """Adds the role + files info box."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = "Light Grid Accent 1"
    cell = tbl.cell(0, 0)
    p = cell.paragraphs[0]
    r = p.add_run("AI Role: ")
    r.bold = True
    r2 = p.add_run(role_text)
    r2.font.size = Pt(11)

    cell = tbl.cell(1, 0)
    p = cell.paragraphs[0]
    r = p.add_run("Primary Files: ")
    r.bold = True
    r2 = p.add_run(files_text)
    r2.font.size = Pt(11)
    r2.italic = True
    doc.add_paragraph()


def add_paragraph_text(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Cm(0.6 * (level + 1))
    run = p.runs[0] if p.runs else p.add_run("")
    p.runs[0].text = text
    return p


def add_qa(q_num, question, answer_paragraphs, code_refs=None):
    """Add a Q&A block with question, multi-paragraph answer, and code refs."""
    # Question
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Q{q_num}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(26, 60, 110)
    run.font.size = Pt(12)
    run2 = p.add_run(question)
    run2.bold = True
    run2.font.size = Pt(12)

    # Answer label + paragraphs
    for i, para in enumerate(answer_paragraphs):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        if i == 0:
            r = p.add_run("Answer: ")
            r.bold = True
            r.font.color.rgb = RGBColor(232, 82, 58)
        p.add_run(para)

    # Code references
    if code_refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Code reference: ")
        r.bold = True
        r.italic = True
        r.font.color.rgb = RGBColor(100, 100, 100)
        for cref in code_refs:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.0)
            p2.paragraph_format.space_after = Pt(2)
            r2 = p2.add_run("• " + cref)
            r2.font.name = "Consolas"
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(60, 60, 60)


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(40, 60, 100)
    # Light background shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F2F8")
    pPr.append(shd)


# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════

add_title("COS40007 — AI for Engineering")
add_subtitle("Lecturer Interview Preparation Guide")
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Project: Smart Government")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(26, 60, 110)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Team: E's Island")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(80, 80, 80)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Team members table
tbl = doc.add_table(rows=5, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "ID"
hdr[1].text = "Name"
hdr[2].text = "Role"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True

members = [
    ("102782494", "Kenneth Hui Hong CHUA", "Inflation Forecasting Lead (ARIMA)"),
    ("102782287", "Kelvin Wen Kiong FONG", "State Grouping Lead (Clustering)"),
    ("102779140", "Nigel Zi Jun LING", "Cross-Analysis & Validation"),
    ("104386364", "Neng Yi CHIENG", "Forecasting Co-Lead + System & Dashboard"),
]
for i, (sid, name, role) in enumerate(members, 1):
    row = tbl.rows[i].cells
    row[0].text = sid
    row[1].text = name
    row[2].text = role

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("How to use this document")
r.bold = True
r.font.size = Pt(13)

intro = (
    "This document is split into four sections, one per team member, plus a final "
    "general project section. Each personal section contains the role description, "
    "in-depth question-and-answer pairs that the lecturer may ask during the "
    "interview, and specific source code line numbers to point to when explaining "
    "your contribution. Answers are written conversationally so they can be "
    "spoken aloud naturally — do not memorise them word-for-word; instead "
    "understand the underlying ideas so you can adapt to follow-up questions."
)
p = doc.add_paragraph(intro)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.right_indent = Cm(1)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1: KENNETH
# ═════════════════════════════════════════════════════════════════════════════

add_section_header("Section 1 — Kenneth Hui Hong CHUA (102782494)")
add_subheader("Role: Inflation Forecasting Lead (ARIMA Model)", color=(232, 82, 58), size=14)

add_role_box(
    "Inflation Forecasting Lead — built the full ARIMA pipeline for forecasting "
    "Malaysia's national CPI inflation, including stationarity testing, automated "
    "order selection, model fitting, fair evaluation against a baseline, and a "
    "fuel-price predictor experiment.",
    "src/preprocessing.py (national CPI loaders), src/time_series.py (ARIMA, "
    "metrics, fuel exogenous experiment)"
)

add_subheader("What Kenneth Did", size=12)
for b in [
    "Collected and cleaned Malaysia's national inflation data from OpenDOSM — "
    "approximately 7,770 monthly records across 14 spending categories from 1980 to 2026.",
    "Built the ARIMA inflation forecasting model with an automated grid search "
    "that tests many parameter combinations and selects the best one by BIC.",
    "Ran the Augmented Dickey-Fuller (ADF) test to confirm the data was stable "
    "enough to forecast reliably.",
    "Trained the final model and justified the chosen settings using standard "
    "information criteria (AIC, BIC).",
    "Measured forecast accuracy using fair metrics — MAE, RMSE, sMAPE, and MASE — "
    "and compared the model against a naive 'next month equals this month' baseline.",
    "Tested whether fuel prices help predict inflation using an ARIMAX model with "
    "fuel as an exogenous regressor.",
]:
    doc.add_paragraph(b, style="List Bullet")

add_subheader("Likely Questions & Detailed Answers", size=13)

add_qa(
    1,
    "Can you explain what ARIMA is and why you chose it over other models?",
    [
        "ARIMA stands for AutoRegressive Integrated Moving Average. It is a classical "
        "statistical model for time series with three components. The AR — autoregressive "
        "part — uses past values of the series itself as predictors. The I — integrated "
        "part — applies differencing to remove trends and make the series stationary. "
        "The MA — moving average part — uses past forecast errors to correct future "
        "predictions. Each component has an order: p for AR, d for differencing, "
        "and q for MA, so the model is written as ARIMA(p, d, q).",

        "We chose ARIMA for three reasons. First, our target is a univariate monthly "
        "series — Malaysia's overall CPI inflation — with over 40 years of history. "
        "That is exactly the kind of data ARIMA is designed for. Second, ARIMA models "
        "are interpretable: every coefficient has a clear statistical meaning, which "
        "matters for a government use case where decisions must be defensible. Third, "
        "we did not have enough labelled future data to train a neural network like "
        "LSTM reliably, and deep learning models would be overkill for a smooth monthly "
        "macroeconomic series.",
    ],
    None,
)

add_qa(
    2,
    "What is the Augmented Dickey-Fuller test and why did you run it?",
    [
        "The ADF test checks whether a time series is stationary. A stationary series "
        "has a constant mean and variance over time and no underlying trend. ARIMA "
        "fundamentally requires stationarity because its mathematical assumptions only "
        "hold when the statistical properties of the data do not drift over time.",

        "The test produces a p-value. If the p-value is below 0.05, we reject the "
        "null hypothesis that the series has a unit root and treat the data as "
        "stationary, so we set d equal to zero. If the p-value is above 0.05, the "
        "series is non-stationary, so we set d to one — meaning we differ the series "
        "once before fitting. This automated decision is wired directly into our order "
        "selection routine so the model adapts to whatever data it receives.",
    ],
    [
        "src/time_series.py lines 38–44 — check_stationarity calls adfuller from "
        "statsmodels and returns the ADF statistic, p-value, and a boolean flag.",
        "src/time_series.py lines 53–54 — inside select_order, d is set to 0 if "
        "the series is stationary, else 1.",
    ],
)

add_qa(
    3,
    "How did you find the best ARIMA order? Why use BIC and not AIC?",
    [
        "We ran a grid search across every combination of p in the range 0 to 4 and "
        "q in the range 0 to 4, giving 24 candidate models in total (we skipped the "
        "degenerate p equals 0 and q equals 0 case because it would just predict the "
        "mean of the differenced series). For each candidate we fitted the ARIMA model "
        "on the training data and recorded its BIC value. We then selected the "
        "(p, d, q) combination with the lowest BIC as our final order.",

        "We preferred BIC — the Bayesian Information Criterion — over AIC because "
        "BIC penalises model complexity more heavily as the sample size grows. "
        "With over 500 months of data, AIC tends to pick overly complex models that "
        "fit the training data closely but generalise poorly to new data. BIC pushes "
        "us towards more parsimonious models, which is what we want for forecasting "
        "two years into the future.",
    ],
    [
        "src/time_series.py lines 47–75 — select_order: BIC grid search.",
        "Lines 61–62 — nested for-loops over p and q.",
        "Line 68 — fits ARIMA(p, d, q) on the training series.",
        "Lines 69–71 — tracks the (p, d, q) tuple with the lowest BIC.",
    ],
)

add_qa(
    4,
    "Walk me through your evaluation metrics. Why these four?",
    [
        "We used four metrics, each chosen for a specific reason.",

        "MAE — Mean Absolute Error — is the simple average of the absolute differences "
        "between our predictions and the actual values. It is in the same units as the "
        "target (percentage points of inflation), so it is easy to interpret.",

        "RMSE — Root Mean Square Error — also returns errors in the original units, "
        "but it penalises large errors more heavily because it squares them first. "
        "A model with one huge mistake will have a much worse RMSE than one with "
        "consistent small mistakes, even if their MAE is identical.",

        "sMAPE — Symmetric Mean Absolute Percentage Error — gives the error as a "
        "percentage. We deliberately avoided regular MAPE because Malaysian YoY "
        "inflation regularly crosses zero — for example, deflation events during 2020. "
        "Regular MAPE divides by the actual value, so it explodes to infinity when "
        "inflation is near zero. sMAPE normalises by the sum of the magnitudes of "
        "actual and predicted, so it stays bounded between 0 and 200 percent.",

        "MASE — Mean Absolute Scaled Error — is the most important one. It divides "
        "our model's MAE by the in-sample MAE of a naive one-step random-walk forecast "
        "on the training data. A MASE below 1 means our ARIMA model beats the baseline "
        "of just repeating last month's value. A MASE above 1 means our model is worse "
        "than doing nothing intelligent. This makes MASE scale-free, comparable across "
        "different datasets, and immune to the zero-crossing problem.",
    ],
    [
        "src/time_series.py lines 92–104 — _smape: handles the zero-crossing case "
        "by masking out points where the denominator is zero.",
        "src/time_series.py lines 107–119 — _mase: computes naive_mae from the "
        "training series via np.diff, then scales the test MAE by it.",
        "src/time_series.py lines 122–137 — evaluate_forecast: assembles all four "
        "metrics into a dictionary.",
    ],
)

add_qa(
    5,
    "What is your naive baseline and why does every forecasting project need one?",
    [
        "Our naive baseline is the random-walk forecast: we predict that the next "
        "24 months will all equal the most recent observed value. It is implemented "
        "in three lines of code — literally np.repeat of the last training value.",

        "Every forecasting project needs a baseline because metrics like RMSE are "
        "meaningless in isolation. An RMSE of 1.5 might be excellent for one dataset "
        "and terrible for another. By comparing our ARIMA's RMSE against the naive "
        "baseline's RMSE, we get a concrete answer to the question 'is our complex "
        "model actually adding value?' If ARIMA cannot beat the naive baseline, "
        "there is no point deploying it. In our case, ARIMA beats the naive baseline, "
        "which means MASE comes out below 1.",
    ],
    [
        "src/time_series.py lines 140–142 — naive_forecast: returns np.repeat of the "
        "last training value for the requested number of steps.",
        "src/time_series.py lines 553–554 in run_all — computes the naive prediction "
        "and evaluates it with the same metrics for a like-for-like comparison.",
    ],
)

add_qa(
    6,
    "You ran an ARIMAX experiment with fuel as an exogenous variable. What did you find and why was the headline forecast still univariate?",
    [
        "The ARIMAX experiment tests whether fuel prices provide predictive signal "
        "beyond what ARIMA already extracts from inflation's own history. We fit two "
        "models on the same fuel-CPI overlap window: an ARIMAX, which uses the average "
        "monthly fuel price as an exogenous regressor, and a plain univariate ARIMA "
        "with no fuel input. Both use the same order so the comparison is fair. We "
        "then compare their RMSE on a held-out 24-month test period.",

        "The result tells us how much fuel improves short-horizon prediction. We also "
        "extracted the fuel coefficient and its p-value from a full-window refit to "
        "check statistical significance.",

        "Despite the explanatory power, the headline two-year forecast deliberately "
        "remains univariate. The honest reason is that ARIMAX needs future values of "
        "the exogenous variable to generate forecasts, and we cannot know future fuel "
        "prices — that is its own forecasting problem. Using ARIMAX for the headline "
        "forecast would require us to forecast fuel prices first and then feed those "
        "uncertain forecasts into the inflation model, compounding the error. "
        "The honest trade-off is: ARIMAX is good for explanation, ARIMA is good for "
        "operational forecasting.",
    ],
    [
        "src/time_series.py lines 173–265 — compute_fuel_exog_model: full ARIMAX vs "
        "ARIMA comparison logic.",
        "Lines 183–186 — the inline comment that documents this explanatory-vs-"
        "operational trade-off.",
        "Lines 225–229 — fits ARIMAX with X_train as the exogenous matrix and "
        "evaluates on X_test.",
        "Lines 242–248 — refits on the full series to extract fuel coefficients and "
        "p-values.",
    ],
)

add_qa(
    7,
    "Explain the data cleaning pipeline for the national CPI dataset.",
    [
        "The raw OpenDOSM file cpi_2d_inflation.csv contains monthly records for "
        "13 spending divisions like Food, Transport, and Health, plus an 'overall' "
        "row, all going back to 1980. We did the following cleaning steps.",

        "First, we parsed the date column as a pandas datetime so we could do time-"
        "based operations. Second, we converted the inflation_yoy and inflation_mom "
        "columns to numeric, telling pandas to coerce any non-numeric entries to NaN "
        "rather than raising an exception. Third, we mapped the two-digit division "
        "codes like '01' to human-readable labels like 'Food and Non-Alcoholic "
        "Beverages' using a hard-coded dictionary.",

        "For the ARIMA model specifically, we extracted only the rows where division "
        "equals 'overall' and dropped any rows with NaN in inflation_yoy. This handles "
        "the start of the series in 1980 where the year-on-year value is undefined "
        "because there is no prior year to compare against.",
    ],
    [
        "src/preprocessing.py lines 21–36 — DIVISION_LABELS dictionary.",
        "src/preprocessing.py lines 39–49 — load_cpi_inflation function.",
        "src/preprocessing.py lines 93–100 — build_overall_inflation extracts the "
        "ARIMA-ready series.",
    ],
)

add_qa(
    8,
    "Why is the test set the last 24 months and not a random 20% of the data?",
    [
        "Time series data has temporal order, and any model evaluation must respect "
        "that order. If we randomly sampled 20% of months for the test set, the model "
        "would see months from 2024 during training and then be asked to predict "
        "months from 2018 — this is called data leakage from the future, and it makes "
        "the evaluation meaningless because real forecasting only ever has access to "
        "the past.",

        "We held out the last 24 months chronologically because it simulates the real "
        "deployment scenario: we have all data up to a cut-off date and need to "
        "forecast the months that follow. We chose 24 months specifically because it "
        "is long enough to capture seasonal patterns and short-term economic cycles, "
        "while still leaving over 500 months of training data. It also matches our "
        "forecast horizon, so the evaluation directly tells us how well the model "
        "would have performed if deployed two years ago.",
    ],
    [
        "src/time_series.py lines 88–89 — train_test_split_ts uses .iloc slicing on "
        "the series with no shuffling.",
        "src/time_series.py line 34 — TEST_MONTHS = 24 constant.",
    ],
)

add_qa(
    9,
    "How does your model decide when to skip the BIC grid search? What is the caching logic?",
    [
        "The BIC grid search is expensive — it fits up to 24 ARIMA models, which can "
        "take several minutes on the full series. We cache the previously found order "
        "in outputs/models/arima_results.json. On the next run, if the file exists and "
        "contains a valid three-element tuple, we reuse that order and skip the grid "
        "search entirely. Deleting the JSON file forces a fresh search, which is what "
        "you do if the input data has changed substantially.",

        "We extend this idea to skip the model refit too: if the cached order matches "
        "AND the current series length matches the cached series length, we know "
        "nothing has changed and we can reuse the train/test metrics, the forecast, "
        "and the pre-fitted model that was saved as a pickle file. This makes "
        "re-running the pipeline almost instant when the data has not changed.",
    ],
    [
        "src/time_series.py lines 497–516 — the cache read logic in run_all.",
        "Lines 519–547 — the broader cache-reuse decision when series length matches.",
        "Lines 568–572 — the freshly fitted full-series model is saved to "
        "arima_fitted.pkl for future reuse.",
    ],
)

add_qa(
    10,
    "What would you do differently if you had more time?",
    [
        "Three things. First, I would extend ARIMA to SARIMA — Seasonal ARIMA — to "
        "explicitly model annual cycles. Malaysian inflation has seasonal patterns "
        "tied to school holidays, festive periods like Hari Raya, and end-of-year "
        "consumption. A seasonal term would let the model capture these without "
        "having to express them through pure autoregression.",

        "Second, I would use a rolling-origin evaluation instead of a single train/test "
        "split. This means re-fitting the model multiple times, each time on a "
        "different chronological window, and averaging the errors. It gives a much "
        "more robust estimate of generalisation than a single 24-month test.",

        "Third, I would explore a hybrid forecasting approach where the long-horizon "
        "univariate ARIMA is combined with a short-horizon ARIMAX for the first three "
        "or six months, where future fuel prices can reasonably be assumed close to "
        "the latest observed level.",
    ],
    None,
)


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2: KELVIN
# ═════════════════════════════════════════════════════════════════════════════

add_section_header("Section 2 — Kelvin Wen Kiong FONG (102782287)")
add_subheader("Role: State Grouping Lead (Clustering)", color=(232, 82, 58), size=14)

add_role_box(
    "State Grouping Lead — engineered five socioeconomic features per state, "
    "standardised them, built both K-Means and Hierarchical clustering, selected "
    "the optimal number of groups using three metrics, validated results with "
    "quality scores, and produced all clustering visualisations.",
    "src/preprocessing.py (state CPI + income loaders, feature engineering), "
    "src/clustering.py (all clustering logic and charts)"
)

add_subheader("What Kelvin Did", size=12)
for b in [
    "Collected and cleaned price and household-income data for all 16 Malaysian states.",
    "Engineered five descriptive measures per state: price level, price growth, "
    "price stability, mean income, and median income.",
    "Standardised all features to a common scale so they could be compared fairly.",
    "Built two different grouping methods (K-Means and Ward Hierarchical) and used "
    "three checks to decide the best number of groups.",
    "Reduced data to 2D using PCA for visualisation, then assigned each group a "
    "clear, meaningful semantic name based on income ordering.",
    "Validated groupings with three quality scores to confirm clusters were genuine "
    "and consistent across both methods.",
    "Produced all clustering charts including elbow plots, PCA scatters, cluster "
    "profile heatmaps, and dendrograms.",
]:
    doc.add_paragraph(b, style="List Bullet")

add_subheader("Likely Questions & Detailed Answers", size=13)

add_qa(
    1,
    "What is clustering and what specific question are you answering for the states?",
    [
        "Clustering is an unsupervised machine learning technique that groups data "
        "points so that points within the same group are more similar to each other "
        "than to points in other groups. Unlike classification, there are no labels — "
        "the algorithm discovers structure in the data on its own.",

        "Our specific question is: which Malaysian states share similar cost-of-"
        "living profiles? Instead of treating all 16 states individually or lumping "
        "them all together, clustering reveals natural socioeconomic tiers — for "
        "example, high-income urban states with high CPI versus low-income rural "
        "states with lower CPI. This is directly useful for government policy because "
        "interventions can be designed for groups rather than crafted state by state.",
    ],
    None,
)

add_qa(
    2,
    "Walk me through the five features you engineered. Why those five?",
    [
        "We engineered five features that together capture both the price environment "
        "and the income environment of each state.",

        "First, mean_cpi_index — the average overall CPI index across all available "
        "months. This captures the absolute price level of the state. A higher value "
        "means goods generally cost more there.",

        "Second, cpi_growth_rate — the percentage change in the CPI index from the "
        "earliest to the latest month. This measures how much prices have grown over "
        "the observation window.",

        "Third, cpi_volatility — the standard deviation of monthly CPI. This is a "
        "proxy for price stability. A state with low volatility has predictable prices; "
        "a state with high volatility experiences shocks.",

        "Fourth, income_mean_latest — the most recent survey mean household income. "
        "This represents purchasing power but is sensitive to high earners.",

        "Fifth, income_median_latest — the most recent median income. This is more "
        "representative of the typical household because medians are not skewed by "
        "outliers.",

        "We included both mean and median income deliberately. If a state has a few "
        "very wealthy households, mean income may misrepresent the typical resident's "
        "purchasing power, but median income tells the truer story. Including both "
        "lets the algorithm pick up on income inequality patterns automatically.",
    ],
    [
        "src/preprocessing.py lines 103–140 — build_state_features assembles all "
        "five features into one row per state.",
        "src/clustering.py lines 35–41 — FEATURE_COLS list defines what the "
        "clustering algorithm sees.",
    ],
)

add_qa(
    3,
    "Why must you standardise features before clustering? What happens if you don't?",
    [
        "Clustering algorithms compute distances between data points. K-Means uses "
        "Euclidean distance; Ward linkage uses variance. Distance calculations are "
        "scale-sensitive: a feature with large numerical values will dominate the "
        "computation simply because of its scale, not because of its informational "
        "importance.",

        "Look at our features. CPI index values are around 100 to 150. Mean income "
        "is in the thousands of Ringgit Malaysia — say 5,000 to 12,000. Growth rate "
        "is a percentage, maybe 20 to 40 percent. Without standardisation, income "
        "would utterly dominate the distance calculation. The algorithm would "
        "essentially be clustering by income alone, and the four other features would "
        "make almost no difference.",

        "We applied StandardScaler, which transforms each feature to have zero mean "
        "and unit variance. After scaling, all five features contribute equally to "
        "the distance metric, and the algorithm can find structure that depends on "
        "the combination of all five.",
    ],
    [
        "src/clustering.py lines 45–49 — prepare_features applies StandardScaler "
        "from sklearn.preprocessing and returns both the scaled array and the fitted "
        "scaler (so it can be inverted later if needed).",
    ],
)

add_qa(
    4,
    "How did you choose the optimal number of clusters? You mentioned three checks.",
    [
        "We tested k values from 2 to 8. Why 8 as the upper limit? With only 16 "
        "states, having k equal to 8 already means each cluster averages just two "
        "members, which is the practical lower bound for a meaningful group. Going "
        "higher would produce singleton clusters.",

        "For each k, we computed three quality metrics. The first is the Silhouette "
        "score, which ranges from -1 to +1. A score near +1 means points are tightly "
        "packed within their own cluster and far from other clusters. A score near "
        "zero means clusters overlap. We treated Silhouette as the primary criterion "
        "and selected the k that maximised it.",

        "The second is the Davies-Bouldin index, where lower is better. It measures "
        "the average ratio of within-cluster scatter to between-cluster separation. "
        "A small value means clusters are compact and well separated.",

        "The third is the Calinski-Harabasz score, where higher is better. It's the "
        "ratio of between-cluster dispersion to within-cluster dispersion.",

        "We also tracked Inertia — the within-cluster sum of squares — for the "
        "classical Elbow Method, where you look for the 'kink' point where adding "
        "more clusters stops dramatically reducing inertia.",
    ],
    [
        "src/clustering.py lines 52–72 — elbow_and_silhouette: the grid search "
        "computing all four metrics for k from 2 to 8.",
        "src/clustering.py line 64 — best_k is selected as the k that maximises "
        "silhouette: np.argmax(silhouettes).",
        "src/clustering.py line 42 — MAX_K constant explicitly set to 8.",
    ],
)

add_qa(
    5,
    "Why did you use K-Means AND Hierarchical clustering? Is one not enough?",
    [
        "Each method makes different mathematical assumptions, so using both gives "
        "us a robustness check.",

        "K-Means assumes clusters are roughly spherical and similar in size. It "
        "requires you to specify k upfront and is computationally efficient. It works "
        "well when the underlying groups have well-defined centres.",

        "Hierarchical clustering with Ward linkage builds a tree of merges by "
        "iteratively combining the two clusters whose union produces the smallest "
        "increase in total within-cluster variance. It does not assume any particular "
        "cluster shape, and it produces a dendrogram that visualises the full hierarchy "
        "of merges. You can decide k after seeing the dendrogram, by cutting it at "
        "different heights.",

        "If both methods produce similar cluster assignments for the same data, we "
        "gain confidence that the structure we found is real and not an artefact of "
        "either algorithm's assumptions. This is essentially methodological "
        "triangulation. In our results, both methods produce consistent clusterings, "
        "which validates the conclusion.",
    ],
    [
        "src/clustering.py lines 75–78 — fit_kmeans with random_state=42 and "
        "n_init=10 for reproducibility.",
        "src/clustering.py lines 81–84 — fit_hierarchical using AgglomerativeClustering "
        "with linkage='ward'.",
    ],
)

add_qa(
    6,
    "Why Ward linkage specifically? There are other linkage methods.",
    [
        "There are several linkage methods in hierarchical clustering. Single linkage "
        "merges based on the closest pair of points between clusters, which creates "
        "long, chain-like clusters and is sensitive to noise. Complete linkage uses "
        "the farthest pair, which tends to produce tight, compact clusters but can "
        "ignore natural elongated structure. Average linkage falls in between.",

        "Ward linkage works differently — it merges the pair of clusters whose union "
        "minimises the increase in total within-cluster variance. This produces "
        "compact, roughly spherical, and relatively balanced clusters. For "
        "socioeconomic data where we expect groups of similar size and where balanced "
        "groupings make policy sense, Ward is empirically the best-performing choice. "
        "It also pairs naturally with K-Means because both methods minimise within-"
        "cluster variance, so they tend to converge on similar solutions.",
    ],
    [
        "src/clustering.py line 82 — AgglomerativeClustering(n_clusters=k, "
        "linkage='ward').",
        "src/clustering.py line 229 — the dendrogram uses Ward linkage too "
        "(scipy.cluster.hierarchy.linkage with method='ward').",
    ],
)

add_qa(
    7,
    "Explain PCA and why you used it. The original data was already 5-dimensional — why reduce further?",
    [
        "Principal Component Analysis is a dimensionality reduction technique. It "
        "finds new axes (called principal components) that are linear combinations of "
        "the original features and chosen so that the first axis captures as much "
        "variance in the data as possible, the second axis captures as much remaining "
        "variance as possible while being orthogonal to the first, and so on.",

        "We used PCA for visualisation only — not for clustering itself. The "
        "clustering algorithm sees all five original standardised features. But "
        "humans cannot visualise 5-D space. By projecting onto the first two "
        "principal components, we get a 2D scatter plot where each state is a point. "
        "The plot's axes are PC1 and PC2, and we label them with the percentage of "
        "variance they explain so the viewer knows how faithful the projection is.",

        "Typically PC1 and PC2 together explain 70 to 90 percent of the variance for "
        "our data, which means the 2D view is a reasonable approximation of the full "
        "5D structure. Clusters that look well-separated in the 2D plot really are "
        "well-separated in 5D, and clusters that look mixed in 2D are often a sign "
        "that the projection is losing important information on PC3 or beyond.",
    ],
    [
        "src/clustering.py lines 87–91 — pca_transform: fits PCA with n_components=2 "
        "and random_state=42.",
        "src/clustering.py lines 176–178 — the plot displays variance percentages on "
        "the axis labels.",
    ],
)

add_qa(
    8,
    "Walk me through your semantic labelling. How did you decide which cluster is 'high income'?",
    [
        "After clustering, the algorithm just gives us integer labels like 0, 1, 2, 3 — "
        "which are meaningless to a policy reader. We needed labels like 'High Income "
        "and High CPI' that communicate what each cluster represents.",

        "Our approach is: for each cluster, compute the mean income across all member "
        "states. Then sort the clusters in ascending order of mean income. The cluster "
        "with the lowest mean income is labelled 'Low Income and Low CPI', the next "
        "one up is 'Mid-Low Income and Mid-Low CPI', then 'Mid-High Income and High "
        "CPI', and finally 'High Income and High CPI'.",

        "This labelling works because in Malaysia, state-level CPI and state-level "
        "income are strongly correlated. Higher-income states like Kuala Lumpur "
        "Federal Territory and Selangor have higher price levels because more "
        "purchasing power drives prices up. Lower-income states like Kelantan and "
        "Perlis have lower price levels. The income-based ordering therefore matches "
        "the price-based ordering, which is why we can use combined labels.",
    ],
    [
        "src/clustering.py lines 94–113 — assign_semantic_labels function.",
        "Line 101 — orders clusters by mean income using argsort on the "
        "income_mean_latest column.",
        "Lines 103–108 — the four pre-defined semantic labels assigned in order.",
    ],
)

add_qa(
    9,
    "What does your dendrogram tell you that the K-Means scatter plot does not?",
    [
        "The dendrogram is the tree diagram produced by hierarchical clustering. It "
        "shows every merge step from the bottom (each state as its own cluster) to "
        "the top (one big cluster containing all states). The height of each merge "
        "indicates how dissimilar the two merging clusters were.",

        "What the dendrogram adds beyond a scatter plot is the hierarchy of "
        "similarity. We can see not just which states are in the same cluster, but "
        "WHICH states are most similar within that cluster, and which clusters merge "
        "next to form super-clusters. For example, we can see that two states might "
        "be in the same final cluster, but one of them is actually closer to a state "
        "in a different cluster — which is information the K-Means scatter would not "
        "expose.",

        "Cutting the dendrogram at different heights gives different numbers of "
        "clusters, so it also serves as a visual tool for choosing k. We coloured "
        "the dendrogram with a colour threshold at 70% of the maximum distance to "
        "highlight the natural cluster boundaries.",
    ],
    [
        "src/clustering.py lines 228–241 — plot_dendrogram using scipy's dendrogram "
        "with Ward linkage and a colour threshold at 70% of max distance.",
        "src/clustering.py lines 302–312 — exports dendrogram coordinates as JSON "
        "so the web dashboard can render an interactive hoverable version.",
    ],
)

add_qa(
    10,
    "How do you defend the choice of MAX_K = 8?",
    [
        "MAX_K is the upper bound of our grid search for the best k. We set it to 8 "
        "based on a sample-size rule of thumb: k should not exceed roughly half the "
        "number of data points, because beyond that each cluster has too few members "
        "to be statistically meaningful or policy-useful.",

        "With 16 states, k equal to 8 means each cluster averages just two states. "
        "Going higher — k of 10 or 12, say — would produce singletons or pairs, "
        "which defeats the purpose of clustering. The whole point is to find groups "
        "with shared characteristics that the government can address as a unit. A "
        "cluster with one state is just that state on its own. So 8 is the practical "
        "upper limit, and in practice the silhouette score peaks well before we get "
        "there.",
    ],
    [
        "src/clustering.py line 42 — MAX_K = 8 with the comment '16 states → max "
        "sensible k is ~half the sample size'.",
    ],
)


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3: NIGEL
# ═════════════════════════════════════════════════════════════════════════════

add_section_header("Section 3 — Nigel Zi Jun LING (102779140)")
add_subheader("Role: Cross-Analysis & Validation", color=(232, 82, 58), size=14)

add_role_box(
    "Cross-Analysis & Validation — connected the forecasting and clustering work "
    "by analysing fuel-price relationships with inflation, building comparison "
    "charts across states and categories, and verifying every key figure against "
    "official 2022 government statistics.",
    "src/preprocessing.py (fuel-price loader), src/time_series.py "
    "(fuel correlation), src/evaluation.py (cross-analysis visualisations)"
)

add_subheader("What Nigel Did", size=12)
for b in [
    "Collected and cleaned the weekly fuel-price data from OpenDOSM and converted "
    "it to monthly frequency to align with the inflation data.",
    "Measured how closely fuel prices move with inflation, both in the same month "
    "(contemporaneous) and one month later (lag-1), for RON95, RON97, and Diesel.",
    "Built state-by-state CPI growth comparison charts to highlight regional disparities.",
    "Built a category-by-year inflation heatmap to identify when and where inflation "
    "was hottest across spending divisions.",
    "Created an income-versus-inflation scatter chart coloured by state cluster, "
    "tying the forecasting and grouping results together.",
    "Cross-checked every key figure against the official Department of Statistics "
    "Malaysia (DOSM) 2022 published numbers to validate the pipeline.",
]:
    doc.add_paragraph(b, style="List Bullet")

add_subheader("Likely Questions & Detailed Answers", size=13)

add_qa(
    1,
    "Why did you need to convert the fuel price data from weekly to monthly?",
    [
        "Our CPI inflation data is recorded monthly — each row is one month, one "
        "country-level value. The fuel price data from OpenDOSM is published weekly. "
        "To measure the relationship between fuel prices and inflation, both series "
        "must share the same time frequency, otherwise we cannot align them in a "
        "merge.",

        "We chose to downsample fuel from weekly to monthly rather than upsample "
        "inflation, because aggregating a higher-frequency series to a lower-frequency "
        "series is statistically sound (you just take the mean within each month), "
        "whereas upsampling would require us to invent intra-month inflation values "
        "we do not have.",

        "Specifically, we used pandas' resample with the 'MS' frequency code, which "
        "stands for month-start, and applied the mean aggregation. So every weekly "
        "price in January gets averaged into one January value. We also computed an "
        "avg_fuel column that is the mean across all three fuel types, which is our "
        "summary indicator for the correlation analysis.",
    ],
    [
        "src/preprocessing.py lines 65–79 — load_fuelprice function.",
        "Line 72 — filters to series_type == 'level' to exclude index and percentage "
        "change rows.",
        "Line 77 — df.set_index('date').resample('MS').mean() does the weekly-to-"
        "monthly conversion.",
        "Line 78 — computes avg_fuel as the mean across ron95, ron97, and diesel.",
    ],
)

add_qa(
    2,
    "What is Pearson correlation and what specifically did you compute for fuel and inflation?",
    [
        "Pearson correlation, denoted r, measures the linear relationship between two "
        "variables. It ranges from minus one to plus one. A value of plus one means "
        "perfect positive correlation — when one variable goes up, the other always "
        "goes up by a proportional amount. Minus one means perfect inverse. Zero "
        "means no linear relationship.",

        "We computed two correlations for each fuel type. The first is contemporaneous "
        "correlation: fuel price at month t versus inflation at month t. This tells us "
        "whether fuel and inflation move together in the same month. The second is "
        "lag-1 correlation: fuel at month t minus 1 versus inflation at month t. "
        "This tests whether fuel changes LEAD inflation changes by one month — that "
        "is, whether fuel price movements predict inflation movements with a delay.",

        "We computed these for four fuel measures: RON95 (the most common consumer "
        "petrol), RON97 (premium petrol), Diesel (used by heavy transport), and "
        "avg_fuel (our composite indicator).",
    ],
    [
        "src/time_series.py lines 145–170 — compute_fuel_correlation function.",
        "Line 160 — contemporaneous: merged['inflation_yoy'].corr(merged[col]).",
        "Line 161 — lag-1: inflation from index 1 onwards correlated with fuel "
        "from index 0 to second-last.",
    ],
)

add_qa(
    3,
    "Why specifically lag-1? Why not lag-2 or lag-6?",
    [
        "Lag-1 was chosen for a specific economic reason. When fuel prices rise, "
        "businesses do not immediately raise consumer prices — there is a "
        "pass-through delay while inventories from before the price change get sold, "
        "supply contracts get renegotiated, and businesses observe whether the price "
        "change is permanent. Empirical economic research suggests this transmission "
        "happens within one to two months for transport-related CPI components in "
        "most economies.",

        "We picked lag-1 as the simplest test of this transmission mechanism. If "
        "the lag-1 correlation is notably stronger than the contemporaneous one, "
        "it supports the leading-indicator hypothesis. If they are similar, fuel and "
        "inflation move together with no clear leader. We chose not to add lag-2 or "
        "lag-6 to keep the analysis focused — too many lags would invite p-hacking "
        "concerns where we might pick whichever lag happens to be strongest.",
    ],
    None,
)

add_qa(
    4,
    "Walk me through the fuel-CPI correlation chart you built.",
    [
        "It is a two-panel chart. The left panel is a scatter plot. Each point "
        "represents one month: the x-coordinate is the average fuel price in Ringgit "
        "Malaysia per litre, and the y-coordinate is overall YoY inflation in percent. "
        "I drew an OLS regression line through the points and annotated it with the "
        "Pearson r and p-value. The visual gives an immediate sense of how tight the "
        "relationship is and which direction it goes.",

        "The right panel is a grouped bar chart. The x-axis shows each fuel type — "
        "RON95, RON97, Diesel, and the composite avg_fuel. For each fuel type there "
        "are two bars side by side: the contemporaneous correlation and the lag-1 "
        "correlation. This makes it easy to see whether different fuel types behave "
        "differently and whether the lag-1 effect is consistently stronger or weaker.",
    ],
    [
        "src/evaluation.py lines 49–95 — plot_fuel_cpi_correlation.",
        "Lines 63–66 — uses scipy.stats.linregress to compute slope, intercept, and "
        "the r and p values shown on the chart.",
        "Lines 75–88 — the side-by-side bar chart of contemporaneous vs lag-1 r "
        "for each fuel type.",
    ],
)

add_qa(
    5,
    "Explain your division inflation heatmap. What insights does it surface?",
    [
        "The heatmap visualises inflation across the two dimensions that matter most "
        "to a policymaker: which category of spending, and which year. Rows are the "
        "13 CPI spending divisions — Food, Transport, Health, and so on. Columns are "
        "years from 2010 to the latest available year. Each cell shows the mean "
        "year-on-year inflation rate for that category in that year.",

        "The colour scale is red-yellow-green reversed, centred on zero. Green cells "
        "mean low or negative inflation. Yellow is moderate. Red is high. Each cell "
        "also has the numeric value annotated for precision.",

        "Several insights emerge from this view. The Transport row turns deep red in "
        "years when fuel subsidies were cut — for example, 2014 and 2022. The Food row "
        "shows inflation persistently above the overall rate after the COVID-19 "
        "pandemic. Some categories like Communication run consistently green, "
        "indicating long-term technology-driven deflation. This kind of pattern is "
        "almost impossible to see in a line chart but obvious in a heatmap.",
    ],
    [
        "src/evaluation.py lines 132–154 — plot_division_inflation_heatmap.",
        "Line 141 — pivot_table with rows=division_label, columns=year, values=mean "
        "of inflation_yoy.",
        "Lines 144–145 — sns.heatmap with cmap='RdYlGn_r', centred on 0, with "
        "annotations enabled.",
    ],
)

add_qa(
    6,
    "What does your state CPI comparison bar chart show?",
    [
        "It is a horizontal bar chart with all 16 Malaysian states on the y-axis and "
        "CPI growth rate as a percentage on the x-axis. The growth rate is computed "
        "as the percentage change in each state's overall CPI index from the earliest "
        "month in our data (2010) to the latest available month.",

        "Bars are coloured by whether the state is above or below the national median "
        "growth rate. Below-median states get a blue bar, above-median states get a "
        "red bar. A dashed vertical line marks the median itself, with its value in "
        "the legend.",

        "This makes regional disparities immediately visible. States in the high-"
        "growth half (red) have experienced faster price increases over the decade, "
        "which suggests their residents have faced relatively worse cost-of-living "
        "pressure. States in the low-growth half (blue) have been more stable. This "
        "is direct input for targeted policy — for example, the government might "
        "consider higher cash subsidies for the high-growth states.",
    ],
    [
        "src/evaluation.py lines 98–129 — plot_state_cpi_comparison.",
        "Lines 105–109 — growth inner function computing percent change from "
        "earliest to latest CPI index per state.",
        "Lines 116–117 — colour assignment based on the median threshold.",
    ],
)

add_qa(
    7,
    "Your income-versus-CPI scatter is the chart that ties everything together. Walk me through it.",
    [
        "This is the synthesis chart that bridges the time-series and clustering work. "
        "Each point is one state. The x-axis is the latest mean household income in "
        "Ringgit Malaysia. The y-axis is the CPI growth rate over the decade. The "
        "colour of each point is the cluster label assigned by Kelvin's K-Means model. "
        "State names are annotated next to each point.",

        "Several things happen on this single chart. First, the strong positive "
        "diagonal trend confirms that higher-income states have generally experienced "
        "higher price growth. Second, the colour grouping shows that the K-Means "
        "clusters separate cleanly along this diagonal — the 'High Income and High "
        "CPI' cluster sits in the top right, 'Low Income and Low CPI' sits in the "
        "bottom left, and the mid-tier clusters fall between them. This visually "
        "validates the cluster labels.",

        "Third, outliers stand out. A state that lies far from the trend line — say, "
        "high income but low CPI growth — would be policy-relevant because it might "
        "represent a counter-example that deserves study.",
    ],
    [
        "src/evaluation.py lines 175–201 — plot_income_vs_cpi_scatter.",
        "Lines 178–180 — colour palette assignment from the cluster labels.",
        "Lines 182–190 — scatter plot loop with per-state annotations.",
    ],
)

add_qa(
    8,
    "How did you cross-check your figures against official 2022 numbers?",
    [
        "Validation against an external source is critical because data cleaning "
        "decisions can silently bias results. We took the official DOSM 2022 "
        "statistical report and compared three categories of figures.",

        "First, the headline overall CPI inflation rate for 2022. We computed this "
        "from our cleaned data by averaging the YoY inflation values across the twelve "
        "months of 2022 and checked it matched the published annual figure.",

        "Second, the mean household income figures per state from the 2022 Household "
        "Income and Expenditure Survey. Our state_features.csv contains "
        "income_mean_latest, which should match the official survey figures for each "
        "of the 16 states.",

        "Third, the CPI division weights and category-level inflation rates. Even if "
        "the overall number is correct, individual categories might be wrong. We "
        "verified the year-end YoY rates for at least Food, Transport, and Housing.",

        "If anything had been off by more than rounding error, we would have known "
        "the preprocessing pipeline had a bug — for example, a unit conversion error "
        "or an incorrect division code mapping.",
    ],
    None,
)

add_qa(
    9,
    "What happens when the fuel data and inflation data cover different time ranges?",
    [
        "Inflation data starts in 1980, but fuel data only starts in 2017. We cannot "
        "compute fuel-inflation correlation on the years before 2017 because the fuel "
        "data does not exist.",

        "We handle this with a pandas inner merge on the date column. An inner merge "
        "keeps only the rows where both datasets have data — so the merged DataFrame "
        "only covers the overlap window from 2017 onwards. We then compute the "
        "correlation on this merged DataFrame.",

        "Crucially, we report the overlap window explicitly in the output: the number "
        "of overlapping months, the start date, and the end date. A reader can then "
        "see that the correlation was computed over, say, 90 monthly observations, "
        "not the full inflation history. This protects against the false impression "
        "that we have decades of fuel-inflation data when in fact we only have a few "
        "years.",
    ],
    [
        "src/time_series.py lines 151–156 — pd.merge with how='inner' on the date "
        "column.",
        "src/time_series.py lines 166–169 — the output dict explicitly reports "
        "n_overlap_months, overlap_start, and overlap_end.",
    ],
)

add_qa(
    10,
    "Why analyse RON95, RON97, and Diesel separately instead of just an average?",
    [
        "Different fuel types serve different economic functions, and lumping them "
        "into a single average could hide important asymmetries.",

        "RON95 is the standard consumer petrol used by most private cars. Its price "
        "is heavily regulated and subsidised in Malaysia, so its movements are "
        "policy-driven and affect household budgets directly.",

        "RON97 is premium petrol, used by a smaller fraction of vehicles, and is not "
        "subject to the same subsidy regime. Its price is closer to a free-market "
        "benchmark.",

        "Diesel powers heavy trucks, buses, generators, and a large share of the "
        "logistics chain. Diesel price changes flow through into the cost of moving "
        "goods, which means they have an indirect but pervasive effect on inflation "
        "across many categories — not just transport.",

        "Looking at all three separately lets us see whether each fuel type has its "
        "own distinctive relationship with inflation. If we only used the average, "
        "a strong diesel-driven effect might be diluted by a weak petrol-driven one, "
        "or vice versa.",
    ],
    None,
)


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4: NENG YI
# ═════════════════════════════════════════════════════════════════════════════

add_section_header("Section 4 — Neng Yi CHIENG (104386364)")
add_subheader("Role: Forecasting Co-Lead + System & Dashboard Builder", color=(232, 82, 58), size=14)

add_role_box(
    "Forecasting Co-Lead and System Builder — co-developed the forecasting "
    "process (train/test split, retraining strategy, confidence intervals, all "
    "forecast charts), designed the four-stage automated pipeline, and built the "
    "interactive web dashboard with four tabs and live data backends.",
    "src/time_series.py (forecasting process and charts), run_pipeline.py "
    "(pipeline orchestrator), app/streamlit_app.py (dashboard), "
    "app/dashboard.py (Flask API)"
)

add_subheader("What Neng Yi Did", size=12)
for b in [
    "Implemented the train/test split for time series, the retrain-on-full-data step, "
    "and the 95% confidence interval extraction.",
    "Built the forecast visualisation charts showing historical data, test-period "
    "actual versus predicted, and the forward forecast with uncertainty band.",
    "Produced the 2-year inflation forecast and saved all results in JSON, PKL, and "
    "CSV formats for reuse across the system.",
    "Added category-level trend charts to provide context for the headline forecast.",
    "Designed the full automated pipeline (run_pipeline.py) that runs preprocessing, "
    "forecasting, clustering, and evaluation in the correct order with deterministic results.",
    "Built the interactive Streamlit dashboard with four tabs: Overview, Forecast, "
    "Clustering, and Fuel Analysis.",
    "Built the Flask backend API (dashboard.py) with REST endpoints and Server-Sent "
    "Events for live pipeline streaming.",
    "Set up Git LFS for large file storage and organised the overall project directory "
    "structure.",
]:
    doc.add_paragraph(b, style="List Bullet")

add_subheader("Likely Questions & Detailed Answers", size=13)

add_qa(
    1,
    "Walk me through how train/test split works differently for time series.",
    [
        "In standard machine learning you randomly shuffle data and take 20 percent "
        "for testing. You cannot do that with time series because every observation "
        "has a timestamp, and the algorithm would end up seeing future months during "
        "training and then being asked to predict past months. This is called data "
        "leakage from the future and it produces wildly optimistic accuracy numbers "
        "that do not hold up in deployment.",

        "Our split respects chronology. We sort the series by date, keep every "
        "observation except the last 24 months for training, and use those final "
        "24 months as the held-out test set. This simulates the actual forecasting "
        "scenario: you have all data up to a cut-off and need to predict the months "
        "that follow.",

        "The split is implemented in two lines using pandas iloc slicing: "
        "series.iloc[:-24] for training and series.iloc[-24:] for testing.",
    ],
    [
        "src/time_series.py lines 88–89 — train_test_split_ts function.",
        "src/time_series.py line 34 — TEST_MONTHS = 24 constant defines the split.",
    ],
)

add_qa(
    2,
    "Why do you retrain on the full data after the test evaluation? Doesn't that contaminate the metrics?",
    [
        "It does not contaminate the metrics, because we report the metrics from the "
        "test split BEFORE retraining. The metrics tell us: 'this is how well the "
        "model would have performed if deployed two years ago.' Once we have that "
        "evaluation result locked in, the train/test split has served its purpose.",

        "Then for the real forward forecast — predicting the next 24 months that "
        "have not happened yet — we fit a fresh ARIMA on all available data including "
        "the test months. We do this because more data generally produces a better-"
        "fitted model, and the most recent observations are especially important "
        "for short-horizon forecasts. If we ignored the test months in the final fit, "
        "we would be throwing away two years of valuable recent information.",

        "So the rule is: train/test split for evaluation, full data for the deployed "
        "forecast.",
    ],
    [
        "src/time_series.py lines 568–572 in run_all — the final fit_arima call uses "
        "the full series, and the result is saved to arima_fitted.pkl for reuse.",
        "src/time_series.py lines 575–579 — get_forecast on the full-data model "
        "produces the headline 2-year prediction.",
    ],
)

add_qa(
    3,
    "Explain the 95% confidence interval on your forecast plot. What does it actually mean?",
    [
        "The 95% confidence interval — sometimes called the prediction interval — is "
        "the violet shaded band around the magenta forecast line. It expresses the "
        "model's uncertainty about its own predictions.",

        "Specifically, under ARIMA's statistical assumptions, there is a 95% "
        "probability that the true future inflation value falls inside the band, for "
        "any given forecast month. The interval comes from ARIMA's internal variance "
        "estimates, which combine the residual variance from the fit and the "
        "propagated uncertainty over the forecast horizon.",

        "Two important features of the band. First, it widens as we forecast further "
        "into the future. This is correct behaviour — uncertainty about month 24 "
        "should be greater than uncertainty about month 1 because errors accumulate. "
        "Second, the band is the model's uncertainty given its assumptions are "
        "correct — if the world changes structurally (a new pandemic, a major policy "
        "shock), the model would not know, and the true value could fall outside the "
        "band. That's a known limitation of any statistical model.",
    ],
    [
        "src/time_series.py lines 82–85 — _extract_forecast: calls "
        "fc.conf_int(alpha=0.05) to get the lower and upper bounds at the 95% level.",
        "src/time_series.py lines 407–412 — the violet shaded band is drawn with "
        "ax.fill_between using these bounds.",
    ],
)

add_qa(
    4,
    "Walk me through the forecast plot — what is every element showing?",
    [
        "The chart has two panels side by side.",

        "Left panel — the headline view. The blue line is historical inflation from "
        "the past ten years; I deliberately truncated to ten years to keep the plot "
        "readable, otherwise four decades of data would compress the recent period "
        "into a sliver. The green line with circle markers is the actual test-period "
        "values from the held-out 24 months. The red dashed line with circles is what "
        "our trained model predicted for those same months — the gap between green "
        "and red is the test error visualised. The magenta line with triangles is the "
        "two-year forward forecast, going beyond the latest observed month. The "
        "violet shaded band is the 95% confidence interval around the forecast. A "
        "vertical grey dotted line marks where training ended and testing began, and "
        "a horizontal black dotted line marks the zero-inflation reference.",

        "Right panel — the test-period zoom. Same data as the left panel's test "
        "section, but isolated and rescaled so the lecturer can clearly see how "
        "closely actual and predicted track each other. The title carries the four "
        "evaluation metrics — MAE, RMSE, sMAPE, and MASE — so the visual and the "
        "numbers tell a consistent story.",
    ],
    [
        "src/time_series.py lines 384–446 — plot_forecast function.",
        "Lines 393–423 — left panel with the truncated 10-year history, test "
        "actual/predicted, forecast, and CI band.",
        "Lines 425–440 — right panel with the test-period zoom and metric title.",
    ],
)

add_qa(
    5,
    "What does the category trend chart add? It is not the headline forecast.",
    [
        "The headline forecast is a single national number — overall CPI inflation. "
        "But that single number is the weighted average of 13 different spending "
        "category inflation rates. The category trend chart breaks that aggregate "
        "apart so the audience can see which components are driving the overall "
        "movement.",

        "The chart shows the last five years of YoY inflation for each spending "
        "category, with each category in its own colour and the overall CPI shown "
        "as a thick black line on top for reference. If the overall line is going up "
        "but Food and Transport are screaming higher while Communication is going "
        "down, the chart immediately reveals the story behind the average.",

        "Without this chart, a policymaker looking at the headline forecast would "
        "not know whether to focus on food, fuel, housing, or healthcare. With it, "
        "the source of inflation pressure is obvious.",
    ],
    [
        "src/time_series.py lines 449–477 — plot_division_trends.",
        "Line 451 — five-year cutoff using pd.DateOffset(years=5).",
        "Lines 461–462 — each category gets its own colour from a tab20 palette; "
        "line 465 emphasises the overall CPI as a thick black line.",
    ],
)

add_qa(
    6,
    "Walk me through your automated pipeline. Why is it set up the way it is?",
    [
        "run_pipeline.py is the master script that orchestrates everything. It runs "
        "four stages in strict order: preprocessing, time-series forecasting, "
        "clustering, and evaluation.",

        "Three design decisions matter. First, at the very top I call os.chdir to "
        "change the working directory to the project root. Without this, relative "
        "paths like data/processed/ break when the script is invoked from a different "
        "directory, which would make the pipeline fragile.",

        "Second, I set MPLCONFIGDIR to a pre-built matplotlib font cache committed "
        "with the repo. On cloud platforms like Streamlit Cloud or Render, matplotlib "
        "rebuilds its font cache on first run, which takes two to three minutes. "
        "Pointing at the pre-built cache makes cold starts almost instant.",

        "Third, I pass Python objects directly between stages rather than re-reading "
        "from disk. The preprocess function returns DataFrames; those DataFrames are "
        "passed straight into run_ts and run_clustering. This makes the pipeline both "
        "faster and easier to debug because the data path through memory is explicit.",

        "Finally, the pipeline launches Streamlit at the end unless --no-dashboard "
        "is passed, so a developer running it locally gets the dashboard automatically.",
    ],
    [
        "run_pipeline.py line 17 — os.chdir to the project root.",
        "run_pipeline.py lines 22–25 — sets MPLCONFIGDIR to the pre-built font cache.",
        "run_pipeline.py lines 43–60 — the four stages, with outputs passed as Python "
        "objects.",
        "run_pipeline.py lines 70–81 — conditional Streamlit launch with --no-dashboard "
        "guard.",
    ],
)

add_qa(
    7,
    "You built a Streamlit dashboard AND a Flask backend. Why both?",
    [
        "They serve different purposes.",

        "Streamlit (app/streamlit_app.py) is the production-facing interactive "
        "dashboard. It is the version the lecturer sees on Streamlit Cloud. Streamlit "
        "is fast to build, integrates natively with pandas and Plotly, supports "
        "auto-deploy from a Git push, and lets us write the entire UI in Python "
        "without touching JavaScript. For a data-science project with a small team, "
        "this is the most productive choice.",

        "The Flask backend (app/dashboard.py) is a separate REST API that exposes "
        "every piece of analysis as a JSON endpoint and supports Server-Sent Events "
        "for live progress streaming. We built it during the prototype phase when we "
        "considered a custom HTML/JavaScript frontend. We kept it in the repository "
        "because it allows any external system — for example, a government's existing "
        "analytics platform — to consume our results programmatically without going "
        "through a browser.",

        "The two can coexist because they share no state. Both read the same JSON, "
        "CSV, and PKL output files produced by the pipeline. Either one can be used "
        "independently.",
    ],
    None,
)

add_qa(
    8,
    "Take me through the four tabs of the Streamlit dashboard.",
    [
        "Tab 1 — Overview. Shows the national overall CPI inflation trend over time, "
        "the CPI breakdown by the 13 spending divisions, an explorer where the user "
        "can pick any state to see its CPI history, and a choropleth map of Malaysia "
        "coloured by mean CPI level per state. This tab is the high-level entry point.",

        "Tab 2 — Time Series Forecast. Displays the chosen ARIMA order, the AIC and "
        "BIC values, and the test-period metrics. Renders the full forecast plot with "
        "historical, test, and forward forecast sections plus the 95% confidence band. "
        "There is also a horizon slider that lets the user request a fresh forecast "
        "for a different number of months, with live progress streaming.",

        "Tab 3 — State Clustering. Lets the user toggle between K-Means and "
        "Hierarchical clustering or view them side by side. Includes the PCA scatter, "
        "the elbow and silhouette plots, the cluster profile heatmap, an interactive "
        "dendrogram, and a choropleth map of Malaysia coloured by cluster assignment.",

        "Tab 4 — Fuel Analysis. Shows the fuel price time series for RON95, RON97, "
        "and Diesel, the contemporaneous and lag-1 correlation bar charts with "
        "inflation, the ARIMAX-versus-ARIMA comparison table from Kenneth's "
        "experiment, and the fuel coefficient and p-value from the full-window refit.",
    ],
    None,
)

add_qa(
    9,
    "Explain Server-Sent Events. Why did you use them?",
    [
        "Server-Sent Events, abbreviated SSE, is a browser standard for one-way "
        "push communication from a server to a client over a long-lived HTTP "
        "connection. Once the connection opens, the server can keep sending "
        "messages until it decides to close the channel.",

        "We used SSE for the long-running pipeline operations. When a user clicks "
        "'Run Pipeline' in the dashboard, the analysis can take minutes — fitting "
        "ARIMA, computing clustering metrics, generating figures. Without SSE, the "
        "user would stare at a blank loading screen with no feedback. With SSE, the "
        "Flask backend spawns the pipeline as a subprocess and streams every line of "
        "stdout to the browser as it appears, so the user sees live progress messages "
        "like 'BIC grid 5 of 24: ARIMA(2,1,3)' and a progress bar that updates "
        "in real time.",

        "The Python side uses a generator function (run_live in time_series.py) that "
        "yields progress dictionaries; the Flask route wraps those yields in SSE-"
        "formatted strings and the browser's EventSource API consumes them. A "
        "heartbeat comment is sent every 15 seconds to keep the connection alive "
        "across proxies that drop idle connections.",
    ],
    [
        "src/time_series.py lines 268–381 — run_live generator that yields progress "
        "events.",
        "app/dashboard.py — /api/ts/live and /api/run endpoints wrap these yields "
        "into SSE format with text/event-stream content type.",
    ],
)

add_qa(
    10,
    "How does the caching strategy work across the whole system?",
    [
        "Caching is layered at three levels.",

        "Layer 1 — disk-level result caching. arima_results.json stores the chosen "
        "ARIMA order, the train/test split details, and the test metrics. "
        "arima_fitted.pkl stores the actual fitted full-series ARIMA model. "
        "clustering_results.json stores the clustering output. On the next pipeline "
        "run, if the input series length matches what is in the cache, we skip the "
        "BIC grid search entirely and reuse the fitted model — saving several minutes.",

        "Layer 2 — Streamlit's @st.cache_data decorator. Every loader function in "
        "the dashboard is wrapped with this decorator. The first time the dashboard "
        "loads a CSV, the result is cached in memory. Subsequent calls with the same "
        "arguments return the cached object instantly. This means UI interactions "
        "like changing a dropdown do not re-read files from disk.",

        "Layer 3 — pre-built matplotlib font cache. As discussed earlier, the "
        "outputs/mpl_cache/ directory contains a pre-built font index that we commit "
        "with the repo. This skips a 2–3 minute font scan on cold starts.",

        "When does the cache invalidate? Layer 1 invalidates if the cached order or "
        "series length does not match the current data. Layer 2 invalidates when the "
        "loader's arguments or the loaded files change. Layer 3 never invalidates — "
        "it is just a static asset.",
    ],
    [
        "src/time_series.py lines 497–534 — disk cache read logic in run_all.",
        "src/time_series.py lines 284–300 — equivalent cache logic in the SSE "
        "streaming path.",
        "run_pipeline.py lines 22–25 and app/streamlit_app.py line 25 — both set "
        "MPLCONFIGDIR to the pre-built font cache.",
    ],
)


# ═════════════════════════════════════════════════════════════════════════════
# GENERAL SECTION
# ═════════════════════════════════════════════════════════════════════════════

add_section_header("Section 5 — General Project Questions (Any Member)")
add_subheader("Topics that may come up regardless of who is being interviewed", color=(232, 82, 58), size=14)

add_qa(
    1,
    "What real-world problem does this project address and who would use it?",
    [
        "Malaysia is in the middle of a cost-of-living debate. Inflation has risen "
        "since 2022, fuel subsidies are being restructured under the RON95 "
        "rationalisation policy, and households across different states are feeling "
        "the impact unevenly. Politicians, economists, and ordinary citizens are "
        "all asking the same question: where is inflation headed, and who is being "
        "hit hardest?",

        "Our system addresses this with three concrete deliverables. A two-year-"
        "ahead national inflation forecast with uncertainty bands. A grouping of "
        "all 16 states into socioeconomic tiers based on shared price and income "
        "patterns. And a quantitative analysis of how fuel prices relate to inflation, "
        "which is directly relevant to the subsidy reform debate.",

        "The target user is a government policymaker — for example, an analyst at "
        "the Ministry of Finance or the Department of Statistics — who needs "
        "evidence-based tools to evaluate policy options. Secondary users include "
        "academic researchers, NGOs working on cost-of-living issues, and engaged "
        "citizens who want a transparent view of the data.",
    ],
    None,
)

add_qa(
    2,
    "Where did all your data come from? How can you trust it?",
    [
        "All four datasets are sourced from OpenDOSM, the official open data portal "
        "of the Department of Statistics Malaysia. DOSM is the federal statistical "
        "agency and the legally authoritative source for economic statistics in "
        "Malaysia. The datasets are publicly published, regularly updated, and used "
        "by Bank Negara Malaysia and the World Bank as primary input.",

        "Our four datasets are: cpi_2d_inflation.csv (national monthly CPI by "
        "13 spending divisions, going back to 1980), cpi_2d_state.csv (monthly CPI "
        "by state and division from 2010), fuelprice.csv (weekly retail fuel prices "
        "from 2017), and hh_income_state.csv (household income by state from the "
        "annual survey, 1970 to 2022).",

        "We trust the data because it is the same data used by the government's "
        "own analysts. We also cross-validated our cleaned outputs against the "
        "official 2022 published figures — Nigel led that effort. If our pipeline "
        "had any systematic preprocessing error, the validation would have caught it.",
    ],
    [
        "data/download_datasets.py — the script with the OpenDOSM URLs for all "
        "four datasets.",
    ],
)

add_qa(
    3,
    "How do you ensure the pipeline is reproducible?",
    [
        "Reproducibility comes from four mechanisms.",

        "First, fixed random seeds. KMeans, PCA, and any other stochastic component "
        "is initialised with random_state=42. Running the pipeline twice on the same "
        "data produces bit-identical outputs.",

        "Second, deterministic ordering. The pipeline always runs preprocessing, "
        "time-series, clustering, evaluation in that order. There is no concurrency "
        "or race condition.",

        "Third, version-controlled inputs. The raw OpenDOSM CSVs are stored in "
        "Git LFS so the exact data used for any given run is preserved.",

        "Fourth, explicit working directory and environment setup at the start of "
        "run_pipeline.py. The chdir and MPLCONFIGDIR calls make the script behave "
        "identically regardless of where Python was invoked from.",

        "If anyone clones the repository and runs 'python run_pipeline.py', they "
        "will get the same JSON outputs, the same figures, and the same fitted "
        "model that we did.",
    ],
    None,
)

add_qa(
    4,
    "How did your team divide the work? Was there any overlap?",
    [
        "Yes, the division was deliberate but not rigid. Kenneth owned the ARIMA "
        "forecasting model itself — the order selection, the metrics, the fuel "
        "exogenous experiment. Kelvin owned the entire clustering analysis — "
        "feature engineering, K-Means, Hierarchical, validation. Nigel owned the "
        "cross-analysis layer — the fuel correlation, the cross-cutting charts, "
        "and the validation against official figures. Neng Yi owned the system — "
        "the pipeline orchestration, the forecasting workflow infrastructure "
        "(train/test split, retraining, plotting), and the dashboards.",

        "Overlap was intentional in two places. Neng Yi shares the 'Forecasting "
        "Co-Lead' title with Kenneth because the forecasting workflow — splitting, "
        "retraining, plotting, saving — sits across both the modelling and the "
        "system layers. Nigel's chart code lives in the same evaluation.py file as "
        "shared utility plots, so we coordinated to avoid duplicate functions.",

        "We used Git branches and weekly meetings to manage the integration. The "
        "single pipeline file (run_pipeline.py) acts as the contract: as long as "
        "each module's run_all returns the right shape, everything fits together.",
    ],
    None,
)

add_qa(
    5,
    "What are the limitations of your approach? What should the lecturer know up front?",
    [
        "Three honest limitations.",

        "First, ARIMA is a univariate statistical model. It assumes the future will "
        "look like a continuation of the past. If there is a structural break — a "
        "major fuel subsidy reform, a global recession, a pandemic — the model "
        "cannot anticipate it. The 95% confidence band is conditional on that "
        "assumption holding.",

        "Second, our clustering uses only five features. There are many other "
        "dimensions of state-level wellbeing that could matter — housing costs, "
        "healthcare access, education quality — that we did not include because the "
        "DOSM open data does not have them at state level on a comparable frequency. "
        "So our clusters reflect cost of living and income, not overall quality of "
        "life.",

        "Third, the fuel-inflation correlation is only computed over the 2017-onward "
        "overlap window, because that is when fuel data starts. The relationship may "
        "have been different in earlier eras under different subsidy regimes.",

        "We document all of these limitations in the report. None of them invalidate "
        "the analysis — they just bound its claims.",
    ],
    None,
)

# ═════════════════════════════════════════════════════════════════════════════
# CLOSING
# ═════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_title("Final Tips for the Interview", size=18)
doc.add_paragraph()

tips = [
    "Open the project in your IDE before the interview so you can navigate to "
    "specific line numbers when needed. The line references in this document "
    "point to the exact functions to display.",

    "Do not memorise these answers verbatim. Read each Q&A two or three times so "
    "you understand the reasoning, then explain it in your own words. The lecturer "
    "will detect rote memorisation immediately.",

    "If you do not know the answer to a follow-up question, say so honestly and "
    "offer to reason it through. A confident 'I am not sure but here is how I "
    "would think about it' is far better than guessing.",

    "If the lecturer asks about a teammate's part, you can briefly summarise it "
    "and then say 'that part was led by [name], they can explain the details.' "
    "Do not pretend to own work that was not yours.",

    "Bring a printed copy of this document and the source code, or have them open "
    "on screen. Visual references make explanations clearer.",

    "Practise explaining one or two of the more complex topics — ARIMA order "
    "selection, the silhouette metric, the train/test split logic — out loud "
    "the day before. Speaking is a different skill from understanding.",
]

for tip in tips:
    p = doc.add_paragraph(tip, style="List Number")
    p.paragraph_format.space_after = Pt(8)


doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Good luck — E's Island")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(26, 60, 110)


# Save
import os
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "COS40007_Interview_Preparation.docx")
doc.save(out_path)
print(f"Saved to {out_path}")
